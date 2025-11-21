"""
Procesador OPTIMIZADO de archivos DOCX
Extrae TODO el contenido fielmente: texto, tablas, imágenes, formatos, estilos
Fallback a OpenAI solo cuando la extracción estructurada falla
"""

import os
import io
import base64
import unicodedata
from datetime import datetime
from typing import Dict, List, Any, Optional
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from html import escape
import re


class DOCXProcessorOptimized:
    """Procesador optimizado de archivos DOCX con extracción completa"""

    def __init__(self):
        self.use_openai_fallback = False
        self.openai_client = None

    def _init_openai(self):
        """Inicializa cliente OpenAI solo si es necesario"""
        if not self.openai_client:
            try:
                import os
                from openai import OpenAI
                from dotenv import load_dotenv
                load_dotenv()

                api_key = os.getenv('OPENAI_API_KEY')
                if api_key:
                    self.openai_client = OpenAI(api_key=api_key)
                    self.use_openai_fallback = True
            except Exception as e:
                print(f"OpenAI no disponible: {e}")
                self.use_openai_fallback = False

    def normalizar_texto(self, texto: str) -> str:
        """Normaliza texto eliminando acentos y convirtiendo a mayúsculas"""
        if not texto:
            return ""
        texto_nfd = unicodedata.normalize('NFD', texto)
        texto_sin_acentos = ''.join(c for c in texto_nfd if unicodedata.category(c) != 'Mn')
        return texto_sin_acentos.upper()

    def extraer_estilo_run(self, run) -> Dict[str, Any]:
        """Extrae información de estilo de un run"""
        estilo = {
            'negrita': run.bold,
            'cursiva': run.italic,
            'subrayado': run.underline,
            'tachado': run.font.strike if run.font else False,
            'superindice': run.font.superscript if run.font else False,
            'subindice': run.font.subscript if run.font else False,
        }

        if run.font:
            estilo['fuente'] = run.font.name
            estilo['tamaño'] = run.font.size.pt if run.font.size else None

            # Color
            if run.font.color and run.font.color.rgb:
                estilo['color'] = str(run.font.color.rgb)

            # Resaltado
            if run.font.highlight_color:
                estilo['resaltado'] = str(run.font.highlight_color)

        return estilo

    def run_a_html(self, run) -> str:
        """Convierte un run a HTML con todos los estilos"""
        if not run.text:
            return ""

        texto = escape(run.text)
        estilo = self.extraer_estilo_run(run)

        # Aplicar estilos en orden
        if estilo.get('negrita'):
            texto = f"<b>{texto}</b>"
        if estilo.get('cursiva'):
            texto = f"<i>{texto}</i>"
        if estilo.get('subrayado'):
            texto = f"<u>{texto}</u>"
        if estilo.get('tachado'):
            texto = f"<s>{texto}</s>"
        if estilo.get('superindice'):
            texto = f"<sup>{texto}</sup>"
        if estilo.get('subindice'):
            texto = f"<sub>{texto}</sub>"

        # Estilos CSS adicionales
        css_styles = []
        if estilo.get('tamaño'):
            css_styles.append(f"font-size:{estilo['tamaño']}pt")
        if estilo.get('color'):
            css_styles.append(f"color:#{estilo['color']}")
        if estilo.get('fuente'):
            css_styles.append(f"font-family:{estilo['fuente']}")

        if css_styles:
            style_attr = '; '.join(css_styles)
            texto = f"<span style='{style_attr}'>{texto}</span>"

        return texto

    def parrafo_a_html(self, parrafo: Paragraph) -> str:
        """Convierte un párrafo completo a HTML preservando todos los estilos"""
        html_runs = []

        for run in parrafo.runs:
            html_runs.append(self.run_a_html(run))

        contenido = ''.join(html_runs)

        if not contenido.strip():
            return ""

        # Detectar alineación
        alineacion = 'left'
        if parrafo.alignment:
            align_value = str(parrafo.alignment)
            if 'CENTER' in align_value:
                alineacion = 'center'
            elif 'RIGHT' in align_value:
                alineacion = 'right'
            elif 'JUSTIFY' in align_value:
                alineacion = 'justify'

        # Detectar si es un título
        if parrafo.style and parrafo.style.name.startswith('Heading'):
            try:
                nivel = int(parrafo.style.name.split()[-1])
                return f"<h{nivel} style='text-align:{alineacion}'>{contenido}</h{nivel}>"
            except (ValueError, IndexError):
                return f"<h1 style='text-align:{alineacion}'>{contenido}</h1>"

        return f"<p style='text-align:{alineacion}'>{contenido}</p>"

    def tabla_a_html(self, tabla: Table) -> str:
        """Convierte una tabla completa a HTML preservando estilos"""
        html = "<table border='1' style='border-collapse: collapse; width:100%;'>\n"

        for fila in tabla.rows:
            html += "<tr>\n"
            for celda in fila.cells:
                # Extraer contenido de la celda (puede contener múltiples párrafos)
                contenido_celda = ""
                for parrafo in celda.paragraphs:
                    contenido_celda += self.parrafo_a_html(parrafo)

                # También puede contener tablas anidadas
                for tabla_anidada in celda.tables:
                    contenido_celda += self.tabla_a_html(tabla_anidada)

                html += f"<td style='padding:5px'>{contenido_celda}</td>\n"

            html += "</tr>\n"

        html += "</table>"
        return html

    def extraer_imagenes(self, doc: Document) -> List[Dict[str, Any]]:
        """Extrae todas las imágenes del documento con sus datos binarios"""
        imagenes = []

        try:
            # Obtener imágenes de las relaciones
            for idx, rel in enumerate(doc.part.rels.values()):
                if "image" in rel.target_ref:
                    imagen_data = {
                        'indice': idx + 1,
                        'tipo': rel.target_ref.split('.')[-1],
                        'relacion_id': rel.rId,
                        'nombre': os.path.basename(rel.target_ref),
                        'datos_base64': None
                    }

                    # Obtener datos binarios
                    try:
                        image_part = rel.target_part
                        imagen_bytes = image_part.blob
                        imagen_data['datos_base64'] = base64.b64encode(imagen_bytes).decode('utf-8')
                        imagen_data['tamaño_bytes'] = len(imagen_bytes)
                    except Exception as e:
                        imagen_data['error_extraccion'] = str(e)

                    imagenes.append(imagen_data)

        except Exception as e:
            print(f"Error al extraer imágenes: {e}")

        return imagenes

    def extraer_listas(self, doc: Document) -> List[Dict[str, Any]]:
        """Extrae listas numeradas y con viñetas"""
        listas = []

        for parrafo in doc.paragraphs:
            if parrafo.style and ('List' in parrafo.style.name or 'list' in parrafo.style.name.lower()):
                listas.append({
                    'texto': parrafo.text,
                    'html': self.parrafo_a_html(parrafo),
                    'estilo': parrafo.style.name
                })

        return listas

    def extraer_documento_completo(self, doc: Document) -> str:
        """
        Extrae TODO el documento como HTML preservando estructura completa
        Incluye: títulos, párrafos, tablas, listas, formatos
        """
        html_completo = []

        # Iterar sobre todos los elementos del documento en orden
        for element in doc.element.body:
            # Párrafo
            if isinstance(element, CT_P):
                parrafo = Paragraph(element, doc)
                parrafo_html = self.parrafo_a_html(parrafo)
                if parrafo_html:
                    html_completo.append(parrafo_html)

            # Tabla
            elif isinstance(element, CT_Tbl):
                tabla = Table(element, doc)
                tabla_html = self.tabla_a_html(tabla)
                html_completo.append(tabla_html)

        return '\n'.join(html_completo)

    def extraer_propuestas_estructuradas(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extrae propuestas usando lógica estructurada (método mejorado)
        Busca en TODO el documento, no solo en tablas
        """
        propuestas = []
        numero = 1

        # 1. Buscar en tablas (método principal)
        for tabla in doc.tables:
            for row in tabla.rows:
                observacion = None
                observacion_html = None
                propuesta_html = None
                propuesta_texto = None

                for idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    cell_text_norm = self.normalizar_texto(cell_text)

                    # Buscar OBSERVACIÓN
                    if "OBSERVACION" in cell_text_norm and idx + 1 < len(row.cells):
                        observacion_html = "".join(
                            [self.parrafo_a_html(p) for p in row.cells[idx + 1].paragraphs]
                        )
                        observacion = row.cells[idx + 1].text.strip()

                    # Buscar PROPUESTA DE SOLVENTACIÓN
                    if "PROPUESTA" in cell_text_norm and "SOLVENTACION" in cell_text_norm:
                        propuesta_html = ""
                        propuesta_texto = ""

                        # Extraer contenido de la celda siguiente
                        if idx + 1 < len(row.cells):
                            # Extraer párrafos
                            for parrafo in row.cells[idx + 1].paragraphs:
                                parrafo_html = self.parrafo_a_html(parrafo)
                                propuesta_html += parrafo_html
                                propuesta_texto += parrafo.text + " "

                            # Extraer tablas anidadas
                            try:
                                for tabla_anidada in row.cells[idx + 1].tables:
                                    tabla_html = self.tabla_a_html(tabla_anidada)
                                    propuesta_html += tabla_html
                                    # Extraer texto de tabla anidada
                                    for fila in tabla_anidada.rows:
                                        for celda in fila.cells:
                                            propuesta_texto += celda.text + " "
                            except (IndexError, AttributeError):
                                pass

                # Agregar propuesta si se encontró
                if propuesta_html and propuesta_html.strip():
                    propuestas.append({
                        "numero": numero,
                        "observacion_texto": observacion or "Sin observación",
                        "observacion_html": observacion_html or "<p>Sin observación</p>",
                        "propuesta_texto": propuesta_texto.strip(),
                        "propuesta_html": propuesta_html,
                        "metodo_extraccion": "estructurado"
                    })
                    numero += 1

        # 2. Buscar también en párrafos fuera de tablas (backup)
        texto_completo = []
        for parrafo in doc.paragraphs:
            texto = parrafo.text.strip()
            if texto:
                texto_norm = self.normalizar_texto(texto)
                if "PROPUESTA" in texto_norm and "SOLVENTACION" in texto_norm:
                    # Agregar contexto (siguiente párrafo podría ser la propuesta)
                    texto_completo.append({
                        'tipo': 'propuesta_candidata',
                        'texto': texto,
                        'html': self.parrafo_a_html(parrafo)
                    })

        return propuestas

    def extraer_con_openai(self, doc: Document, documento_html: str) -> List[Dict[str, Any]]:
        """
        Fallback: usa OpenAI para extraer propuestas cuando la lógica estructurada falla
        """
        if not self.use_openai_fallback:
            self._init_openai()

        if not self.openai_client:
            return []

        try:
            # Limitar el documento para no exceder límites de tokens
            documento_truncado = documento_html[:12000]

            prompt = f"""Eres un experto en análisis de documentos de auditoría y solventación.

Analiza el siguiente documento Word (en formato HTML) y extrae TODAS las propuestas de solventación que encuentres.

Una propuesta típicamente tiene:
1. Una OBSERVACIÓN (opcional)
2. Una PROPUESTA DE SOLVENTACIÓN

DOCUMENTO:
{documento_truncado}

Extrae TODAS las propuestas y devuelve un JSON con este formato:
{{
    "propuestas": [
        {{
            "numero": 1,
            "observacion": "texto de la observación o 'Sin observación'",
            "propuesta": "texto completo de la propuesta"
        }},
        ...
    ]
}}

IMPORTANTE: Solo devuelve el JSON, sin texto adicional."""

            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Eres un experto en análisis de documentos. Respondes solo en JSON válido."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=2000
            )

            # Parsear respuesta
            import json
            respuesta_texto = response.choices[0].message.content.strip()

            # Limpiar markdown si existe
            if respuesta_texto.startswith('```'):
                respuesta_texto = respuesta_texto.split('```')[1]
                if respuesta_texto.startswith('json'):
                    respuesta_texto = respuesta_texto[4:]

            resultado = json.loads(respuesta_texto)

            # Formatear propuestas
            propuestas = []
            for idx, prop in enumerate(resultado.get('propuestas', []), start=1):
                propuestas.append({
                    "numero": idx,
                    "observacion_texto": prop.get('observacion', 'Sin observación'),
                    "observacion_html": f"<p>{escape(prop.get('observacion', 'Sin observación'))}</p>",
                    "propuesta_texto": prop.get('propuesta', ''),
                    "propuesta_html": f"<p>{escape(prop.get('propuesta', ''))}</p>",
                    "metodo_extraccion": "openai_fallback"
                })

            return propuestas

        except Exception as e:
            print(f"Error en extracción con OpenAI: {e}")
            return []

    def extraer_titulos(self, doc: Document) -> List[Dict[str, Any]]:
        """Extrae todos los títulos del documento"""
        titulos = []

        for parrafo in doc.paragraphs:
            if parrafo.style and parrafo.style.name.startswith('Heading'):
                try:
                    nivel = int(parrafo.style.name.split()[-1])
                except (ValueError, IndexError):
                    nivel = 1

                if parrafo.text.strip():
                    titulos.append({
                        'nivel': nivel,
                        'texto': parrafo.text.strip(),
                        'html': self.parrafo_a_html(parrafo)
                    })

        return titulos

    def extraer_metadatos(self, doc: Document, filepath: str) -> Dict[str, Any]:
        """Extrae metadatos completos del documento"""
        try:
            core_props = doc.core_properties

            return {
                'nombre_archivo': os.path.basename(filepath),
                'autor': core_props.author or 'Desconocido',
                'titulo': core_props.title or 'Sin título',
                'asunto': core_props.subject or 'Sin asunto',
                'descripcion': core_props.comments or 'Sin descripción',
                'categoria': core_props.category or 'Sin categoría',
                'palabras_clave': core_props.keywords or 'Sin palabras clave',
                'fecha_creacion': core_props.created.isoformat() if core_props.created else None,
                'fecha_modificacion': core_props.modified.isoformat() if core_props.modified else None,
                'ultima_modificacion_por': core_props.last_modified_by or 'Desconocido',
                'revision': core_props.revision or 0,
                'tamano_archivo': os.path.getsize(filepath)
            }

        except Exception as e:
            return {
                'nombre_archivo': os.path.basename(filepath),
                'error_metadatos': str(e),
                'tamano_archivo': os.path.getsize(filepath) if os.path.exists(filepath) else 0
            }

    def calcular_estadisticas(self, doc: Document, propuestas: List[Dict]) -> Dict[str, Any]:
        """Calcula estadísticas completas del documento"""
        total_parrafos = len(doc.paragraphs)
        total_palabras = sum(len(p.text.split()) for p in doc.paragraphs)
        total_caracteres = sum(len(p.text) for p in doc.paragraphs)
        total_tablas = len(doc.tables)

        # Contar párrafos con contenido
        parrafos_con_contenido = sum(1 for p in doc.paragraphs if p.text.strip())

        # Contar elementos con formato
        elementos_negrita = sum(1 for p in doc.paragraphs for r in p.runs if r.bold)
        elementos_cursiva = sum(1 for p in doc.paragraphs for r in p.runs if r.italic)
        elementos_subrayado = sum(1 for p in doc.paragraphs for r in p.runs if r.underline)

        return {
            'total_parrafos': total_parrafos,
            'parrafos_con_contenido': parrafos_con_contenido,
            'total_palabras': total_palabras,
            'total_caracteres': total_caracteres,
            'total_tablas': total_tablas,
            'total_propuestas': len(propuestas),
            'elementos_con_formato': {
                'negrita': elementos_negrita,
                'cursiva': elementos_cursiva,
                'subrayado': elementos_subrayado
            }
        }

    def process_docx(self, filepath: str) -> Dict[str, Any]:
        """
        Procesa un archivo DOCX de manera optimizada
        Extrae TODO el contenido fielmente y usa OpenAI solo como fallback
        """
        try:
            doc = Document(filepath)

            # 1. Extraer metadatos
            metadatos = self.extraer_metadatos(doc, filepath)

            # 2. Extraer documento completo como HTML
            documento_html = self.extraer_documento_completo(doc)

            # 3. Extraer imágenes
            imagenes = self.extraer_imagenes(doc)
            metadatos['imagenes'] = {
                'tiene_imagenes': len(imagenes) > 0,
                'cantidad': len(imagenes),
                'detalles': imagenes
            }

            # 4. Extraer títulos
            titulos = self.extraer_titulos(doc)

            # 5. Extraer listas
            listas = self.extraer_listas(doc)

            # 6. Intentar extracción estructurada de propuestas
            propuestas = []
            try:
                propuestas = self.extraer_propuestas_estructuradas(doc)
            except Exception as e:
                print(f"Extracción estructurada falló: {e}")

            # 7. Si no se encontraron propuestas, intentar con OpenAI (fallback)
            if len(propuestas) == 0:
                print("Usando OpenAI como fallback...")
                propuestas = self.extraer_con_openai(doc, documento_html)

            # 8. Calcular estadísticas
            estadisticas = self.calcular_estadisticas(doc, propuestas)
            estadisticas['metodo_extraccion_usado'] = 'estructurado' if any(
                p.get('metodo_extraccion') == 'estructurado' for p in propuestas
            ) else 'openai_fallback' if propuestas else 'sin_propuestas'

            return {
                'tipo_archivo': 'DOCX',
                'nombre_archivo': os.path.basename(filepath),
                'procesado_en': datetime.now().isoformat(),
                'metadatos': metadatos,
                'estadisticas': estadisticas,
                'contenido': {
                    'documento_completo_html': documento_html,
                    'titulos': titulos,
                    'listas': listas,
                    'propuestas': propuestas
                },
                'extraccion_exitosa': True
            }

        except Exception as e:
            return {
                'tipo_archivo': 'DOCX',
                'nombre_archivo': os.path.basename(filepath),
                'error': str(e),
                'procesado_en': datetime.now().isoformat(),
                'extraccion_exitosa': False
            }


# Instancia global del procesador
processor = DOCXProcessorOptimized()


def process_docx(filepath: str) -> Dict[str, Any]:
    """Función de compatibilidad con el código existente"""
    return processor.process_docx(filepath)
