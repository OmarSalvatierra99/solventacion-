"""
Procesador de archivos DOCX
Extrae propuestas de solventación, metadatos, y detecta imágenes
"""

import os
import unicodedata
from datetime import datetime
from docx import Document
from html import escape


def normalizar_texto(texto):
    """
    Normaliza texto eliminando acentos y convirtiendo a mayúsculas.

    Args:
        texto (str): Texto a normalizar.

    Returns:
        str: Texto normalizado.
    """
    if not texto:
        return ""
    # Eliminar acentos
    texto_nfd = unicodedata.normalize('NFD', texto)
    texto_sin_acentos = ''.join(c for c in texto_nfd if unicodedata.category(c) != 'Mn')
    return texto_sin_acentos.upper()


def convertir_parrafo_a_html_crudo(parrafo):
    """
    Convierte un párrafo de Word a HTML crudo respetando negritas, cursivas y subrayado.

    Args:
        parrafo (docx.text.Paragraph): Párrafo del documento Word.

    Returns:
        str: Representación HTML del párrafo en crudo.
    """
    html = ""
    for run in parrafo.runs:
        texto = escape(run.text)  # Asegurar caracteres HTML válidos
        if run.bold:
            texto = f"<b>{texto}</b>"
        if run.italic:
            texto = f"<i>{texto}</i>"
        if run.underline:
            texto = f"<u>{texto}</u>"
        html += texto
    return f"<p style='text-align:justify'>{html}</p>" if html.strip() else ""


def convertir_tabla_a_html_crudo(tabla):
    """
    Convierte una tabla de Word en formato HTML crudo.

    Args:
        tabla (docx.table.Table): Tabla del documento Word.

    Returns:
        str: Representación HTML cruda de la tabla.
    """
    html = "<table border='1' style='border-collapse: collapse;'>\n"
    for fila in tabla.rows:
        html += "<tr>\n"
        for celda in fila.cells:
            contenido = "".join([convertir_parrafo_a_html_crudo(parrafo) for parrafo in celda.paragraphs])
            html += f"<td>{contenido}</td>\n"
        html += "</tr>\n"
    html += "</table>"
    return html


def detectar_imagenes(doc):
    """
    Detecta si el documento contiene imágenes.

    Args:
        doc (Document): Documento de Word.

    Returns:
        dict: Información sobre imágenes en el documento
    """
    imagenes = []

    # Buscar imágenes en las relaciones del documento
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            imagenes.append({
                'tipo': rel.target_ref.split('.')[-1],
                'relacion_id': rel.rId
            })

    return {
        'tiene_imagenes': len(imagenes) > 0,
        'cantidad': len(imagenes),
        'detalles': imagenes
    }


def extraer_propuestas(doc):
    """
    Extrae propuestas de solventación y observaciones de tablas en el documento.

    Args:
        doc (Document): Documento de Word.

    Returns:
        list[dict]: Lista de propuestas con su número, observación y contenido en HTML.
    """
    datos_extraidos = []
    numero = 1

    for tabla in doc.tables:
        for row in tabla.rows:
            observacion = None
            propuesta_html = None

            for idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                cell_text_norm = normalizar_texto(cell_text)

                # Buscar "OBSERVACIÓN" en la celda (normalizado)
                if "OBSERVACION" in cell_text_norm and idx + 1 < len(row.cells):
                    observacion_contenido = "".join(
                        [convertir_parrafo_a_html_crudo(parrafo) for parrafo in row.cells[idx + 1].paragraphs]
                    )
                    observacion = observacion_contenido if observacion_contenido.strip() else "Sin observación"

                # Buscar "PROPUESTA DE SOLVENTACIÓN" en la celda (normalizado para aceptar cualquier variante)
                if "PROPUESTA" in cell_text_norm and "SOLVENTACION" in cell_text_norm:
                    propuesta_html = ""

                    # Extraer contenido de la celda siguiente (con protección de índice)
                    if idx + 1 < len(row.cells):
                        contenido_html = "".join(
                            [convertir_parrafo_a_html_crudo(parrafo) for parrafo in row.cells[idx + 1].paragraphs]
                        )
                        propuesta_html += contenido_html

                        # Extraer tablas dentro de la celda
                        try:
                            for tabla_asociada in row.cells[idx + 1].tables:
                                tabla_html_crudo = convertir_tabla_a_html_crudo(tabla_asociada)
                                propuesta_html += tabla_html_crudo
                        except (IndexError, AttributeError):
                            pass  # Continuar si no hay tablas o hay error de acceso

            # Si encontramos una propuesta, agregarla a los datos
            if propuesta_html and propuesta_html.strip():
                datos_extraidos.append({
                    "numero": numero,
                    "observacion": observacion or "Sin observación",
                    "propuesta_html": propuesta_html
                })
                numero += 1

    return datos_extraidos


def extraer_metadatos(doc, filepath):
    """
    Extrae metadatos del documento.

    Args:
        doc (Document): Documento de Word.
        filepath (str): Ruta del archivo.

    Returns:
        dict: Metadatos del documento.
    """
    core_props = doc.core_properties

    return {
        'nombre_archivo': os.path.basename(filepath),
        'autor': core_props.author or 'Desconocido',
        'titulo': core_props.title or 'Sin título',
        'asunto': core_props.subject or 'Sin asunto',
        'fecha_creacion': core_props.created.isoformat() if core_props.created else None,
        'fecha_modificacion': core_props.modified.isoformat() if core_props.modified else None,
        'ultima_modificacion_por': core_props.last_modified_by or 'Desconocido',
        'revision': core_props.revision or 0,
        'tamano_archivo': os.path.getsize(filepath)
    }


def calcular_estadisticas(doc, propuestas):
    """
    Calcula estadísticas del documento.

    Args:
        doc (Document): Documento de Word.
        propuestas (list): Lista de propuestas extraídas.

    Returns:
        dict: Estadísticas del documento.
    """
    total_parrafos = len(doc.paragraphs)
    total_palabras = sum(len(p.text.split()) for p in doc.paragraphs)
    total_tablas = len(doc.tables)
    total_propuestas = len(propuestas)

    # Contar párrafos no vacíos
    parrafos_con_contenido = sum(1 for p in doc.paragraphs if p.text.strip())

    return {
        'total_parrafos': total_parrafos,
        'parrafos_con_contenido': parrafos_con_contenido,
        'total_palabras': total_palabras,
        'total_tablas': total_tablas,
        'total_propuestas': total_propuestas
    }


def extraer_titulos(doc):
    """
    Extrae títulos del documento basándose en estilos.

    Args:
        doc (Document): Documento de Word.

    Returns:
        list[dict]: Lista de títulos con su nivel y texto.
    """
    titulos = []

    for parrafo in doc.paragraphs:
        if parrafo.style.name.startswith('Heading'):
            try:
                nivel = int(parrafo.style.name.split()[-1])
            except (ValueError, IndexError):
                nivel = 1

            if parrafo.text.strip():
                titulos.append({
                    'nivel': nivel,
                    'texto': parrafo.text.strip()
                })

    return titulos


def process_docx(filepath):
    """
    Procesa un archivo DOCX y extrae toda la información relevante.

    Args:
        filepath (str): Ruta del archivo DOCX.

    Returns:
        dict: Información extraída del documento.
    """
    try:
        doc = Document(filepath)

        # Extraer propuestas de solventación
        propuestas = extraer_propuestas(doc)

        # Detectar imágenes
        info_imagenes = detectar_imagenes(doc)

        # Extraer metadatos
        metadatos = extraer_metadatos(doc, filepath)
        metadatos['imagenes'] = info_imagenes

        # Calcular estadísticas
        estadisticas = calcular_estadisticas(doc, propuestas)

        # Extraer títulos
        titulos = extraer_titulos(doc)

        return {
            'tipo_archivo': 'DOCX',
            'nombre_archivo': os.path.basename(filepath),
            'procesado_en': datetime.now().isoformat(),
            'metadatos': metadatos,
            'estadisticas': estadisticas,
            'contenido': {
                'titulos': titulos,
                'propuestas': propuestas
            }
        }

    except Exception as e:
        return {
            'tipo_archivo': 'DOCX',
            'nombre_archivo': os.path.basename(filepath),
            'error': str(e),
            'procesado_en': datetime.now().isoformat()
        }
